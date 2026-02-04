#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a small synthetic dataset for reproducibility (no confidential data).

Outputs (project-relative):
- Supplementary/synthetic_dataset/documents_zh/*.docx (Chinese; actual inputs used by the pipeline)
- Supplementary/synthetic_dataset/documents_en/*.docx (English translations for reviewer readability)
- Supplementary/synthetic_dataset/data/survey_points_synthetic.csv
- Supplementary/synthetic_dataset/data/ground_truth_annotations_synthetic.csv
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import csv
import argparse
from datetime import date


@dataclass(frozen=True)
class SyntheticHole:
    hole_id: str
    location_desc: str
    location_desc_en: str
    design_depth_m: float | None
    design_azimuth_deg: float | None
    design_inclination_deg: float | None
    actual_depth_m: float | None
    actual_azimuth_deg: float | None
    actual_inclination_deg: float | None


@dataclass(frozen=True)
class SyntheticDoc:
    filename: str
    title_zh: str
    project_line_zh: str
    title_en: str
    project_line_en: str
    report_date: str
    holes: list[SyntheticHole]


def _write_csv(path: Path, rows: list[dict[str, object]], fieldnames: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def _write_docx_zh(path: Path, doc: SyntheticDoc) -> None:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "python-docx is required to generate .docx files. "
            "Install with: pip install python-docx"
        ) from e

    path.parent.mkdir(parents=True, exist_ok=True)

    d = Document()
    d.add_heading(doc.title_zh, level=1)
    d.add_paragraph(doc.project_line_zh)
    d.add_paragraph(f"报告日期：{doc.report_date}")
    d.add_paragraph("说明：本文件为合成（虚拟）数据，用于复现流程逻辑，不包含任何真实工程项目的保密信息。")
    d.add_paragraph("")

    d.add_paragraph(
        f"本工程于{doc.report_date}完成验收，共施工{len(doc.holes)}个钻孔。现将钻孔验收与施工记录汇总如下："
    )
    d.add_paragraph("")

    d.add_heading("一、任务来源", level=2)
    d.add_paragraph("本次钻探用于示例说明钻孔竣工报告的记录方式，并提供结构化信息抽取与坐标推断所需字段。")
    d.add_paragraph("")

    d.add_heading("二、钻孔施工情况", level=2)

    for idx, hole in enumerate(doc.holes, start=1):
        d.add_paragraph(
            f"{idx}．{hole.hole_id}钻孔，设计位置：{hole.location_desc}；"
            f"设计方位：{hole.design_azimuth_deg}°；设计倾角：{hole.design_inclination_deg}°；"
            f"设计孔深：{hole.design_depth_m} m。"
        )
        d.add_paragraph(
            f"该孔实际施工参数：实际方位{hole.actual_azimuth_deg}°，实际倾角{hole.actual_inclination_deg}°，"
            f"终孔深度{hole.actual_depth_m} m。"
        )

        d.add_paragraph("参数汇总表：")
        t = d.add_table(rows=1, cols=4)
        hdr = t.rows[0].cells
        hdr[0].text = "字段"
        hdr[1].text = "设计"
        hdr[2].text = "实际"
        hdr[3].text = "单位"

        def add_row(field: str, design: str, actual: str, unit: str) -> None:
            r = t.add_row().cells
            r[0].text = field
            r[1].text = design
            r[2].text = actual
            r[3].text = unit

        add_row("孔深", str(hole.design_depth_m), str(hole.actual_depth_m), "m")
        add_row("方位角", str(hole.design_azimuth_deg), str(hole.actual_azimuth_deg), "°")
        add_row("倾角", str(hole.design_inclination_deg), str(hole.actual_inclination_deg), "°")
        d.add_paragraph("")

    d.add_heading("三、结论", level=2)
    d.add_paragraph("以上钻孔记录为合成数据，仅用于复现信息抽取与坐标推断流程，不代表任何真实工程项目。")

    d.save(path)

def _write_docx_en(path: Path, doc: SyntheticDoc) -> None:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "python-docx is required to generate .docx files. "
            "Install with: pip install python-docx"
        ) from e

    path.parent.mkdir(parents=True, exist_ok=True)

    d = Document()
    d.add_heading(doc.title_en, level=1)
    d.add_paragraph(doc.project_line_en)
    d.add_paragraph(f"Report date: {doc.report_date}")
    d.add_paragraph(
        "NOTE: This document is synthetic (dummy) data for reproducibility and contains no confidential content."
    )
    d.add_paragraph("")

    d.add_paragraph(
        f"This completion report summarizes {len(doc.holes)} synthetic boreholes and demonstrates a realistic writing style used in borehole records."
    )
    d.add_paragraph("")

    d.add_heading("1. Task background", level=2)
    d.add_paragraph(
        "This synthetic report is provided to support reproducibility of the information-extraction and coordinate-inference pipeline, without redistributing confidential project documents."
    )
    d.add_paragraph("")

    d.add_heading("2. Borehole records", level=2)
    for idx, hole in enumerate(doc.holes, start=1):
        d.add_paragraph(
            f"{idx}. Borehole {hole.hole_id}. Design location: {hole.location_desc_en}; "
            f"design azimuth: {hole.design_azimuth_deg}°; design inclination: {hole.design_inclination_deg}°; "
            f"design depth: {hole.design_depth_m} m."
        )
        d.add_paragraph(
            f"Actual drilling parameters: azimuth {hole.actual_azimuth_deg}°, inclination {hole.actual_inclination_deg}°, "
            f"final depth {hole.actual_depth_m} m."
        )

        d.add_paragraph("Key parameters table:")
        t = d.add_table(rows=1, cols=4)
        hdr = t.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Design"
        hdr[2].text = "Actual"
        hdr[3].text = "Unit"

        def add_row(field: str, design: str, actual: str, unit: str) -> None:
            r = t.add_row().cells
            r[0].text = field
            r[1].text = design
            r[2].text = actual
            r[3].text = unit

        add_row("Depth", str(hole.design_depth_m), str(hole.actual_depth_m), "m")
        add_row("Azimuth", str(hole.design_azimuth_deg), str(hole.actual_azimuth_deg), "deg")
        add_row("Inclination", str(hole.design_inclination_deg), str(hole.actual_inclination_deg), "deg")
        d.add_paragraph("")

    d.add_heading("3. Notes", level=2)
    d.add_paragraph(
        "All contents are synthetic and for demonstration only; they do not correspond to any real engineering project."
    )
    d.save(path)

def generate(project_root: Path, output_root: Path) -> dict[str, Path]:
    docs_zh_dir = output_root / "documents_zh"
    docs_en_dir = output_root / "documents_en"
    data_dir = output_root / "data"

    docs: list[SyntheticDoc] = [
        SyntheticDoc(
            filename="synthetic_borehole_report_001.docx",
            title_zh="合成钻孔竣工报告 001（虚拟数据）",
            project_line_zh="合成项目：巷道超前钻探（虚拟数据）",
            title_en="Synthetic Borehole Completion Report 001 (Dummy Data)",
            project_line_en="Synthetic project line: roadway advance drilling (dummy data)",
            report_date=str(date(2026, 1, 20)),
            holes=[
                SyntheticHole(
                    hole_id="SYN-303C-01",
                    location_desc="运输巷 316#点前 65m",
                    location_desc_en="65 m before point 316# in the transport roadway",
                    design_depth_m=100.0,
                    design_azimuth_deg=90.0,
                    design_inclination_deg=-10.0,
                    actual_depth_m=98.0,
                    actual_azimuth_deg=92.0,
                    actual_inclination_deg=-11.0,
                ),
                SyntheticHole(
                    hole_id="SYN-303C-02",
                    location_desc="轨道上山 15#点前 88m",
                    location_desc_en="88 m before point 15# on the haulage roadway",
                    design_depth_m=60.0,
                    design_azimuth_deg=45.0,
                    design_inclination_deg=-5.0,
                    actual_depth_m=58.0,
                    actual_azimuth_deg=44.0,
                    actual_inclination_deg=-6.0,
                ),
            ],
        ),
        SyntheticDoc(
            filename="synthetic_borehole_report_002.docx",
            title_zh="合成钻孔竣工报告 002（虚拟数据）",
            project_line_zh="合成项目：巷道水文探查（虚拟数据）",
            title_en="Synthetic Borehole Completion Report 002 (Dummy Data)",
            project_line_en="Synthetic project line: roadway hydrogeological probing (dummy data)",
            report_date=str(date(2026, 1, 21)),
            holes=[
                SyntheticHole(
                    hole_id="SYN-303C-03",
                    location_desc="CP12与CP13之间偏左 5m",
                    location_desc_en="between CP12 and CP13, 5 m to the left",
                    design_depth_m=80.0,
                    design_azimuth_deg=270.0,
                    design_inclination_deg=-12.0,
                    actual_depth_m=79.0,
                    actual_azimuth_deg=268.0,
                    actual_inclination_deg=-12.0,
                ),
                SyntheticHole(
                    hole_id="SYN-303C-04",
                    location_desc="15#点后 30m",
                    location_desc_en="30 m after point 15#",
                    design_depth_m=40.0,
                    design_azimuth_deg=0.0,
                    design_inclination_deg=-3.0,
                    actual_depth_m=39.0,
                    actual_azimuth_deg=2.0,
                    actual_inclination_deg=-4.0,
                ),
            ],
        ),
    ]

    # Survey points (minimal)
    survey_points_file = data_dir / "survey_points_synthetic.csv"
    _write_csv(
        survey_points_file,
        rows=[
            {"FID": "15", "X": 4097000.0, "Y": 40000.0, "Z": -320.0, "name": "synthetic"},
            {"FID": "16", "X": 4097100.0, "Y": 40000.0, "Z": -320.0, "name": "synthetic"},
            {"FID": "CP12", "X": 4097000.0, "Y": 40100.0, "Z": -320.0, "name": "synthetic"},
            {"FID": "CP13", "X": 4097000.0, "Y": 40200.0, "Z": -320.0, "name": "synthetic"},
        ],
        fieldnames=["FID", "X", "Y", "Z", "name"],
    )

    # Ground truth (document-level counts only)
    gt_file = data_dir / "ground_truth_annotations_synthetic.csv"
    _write_csv(
        gt_file,
        rows=[
            {
                "document_filename": d.filename,
                "true_total_entities_count": len(d.holes),
                "true_entities_with_location_count": sum(1 for h in d.holes if h.location_desc),
            }
            for d in docs
        ],
        fieldnames=[
            "document_filename",
            "true_total_entities_count",
            "true_entities_with_location_count",
        ],
    )

    # Docs
    for d in docs:
        # Chinese: actual inputs used by the pipeline
        _write_docx_zh(docs_zh_dir / d.filename, d)

        # English: translations for reviewer readability
        _write_docx_en(docs_en_dir / d.filename, d)

    return {
        "documents_zh_dir": docs_zh_dir,
        "documents_en_dir": docs_en_dir,
        "survey_points_file": survey_points_file,
        "ground_truth_file": gt_file,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a bilingual synthetic dataset (no confidential data).")
    parser.add_argument(
        "--output-root",
        type=Path,
        default=Path("Supplementary/synthetic_dataset"),
        help="Output folder (relative to project root by default).",
    )
    args = parser.parse_args()

    project_root = Path(__file__).resolve().parents[1]
    output_root = (project_root / args.output_root).resolve() if not args.output_root.is_absolute() else args.output_root
    paths = generate(project_root, output_root)

    print("Synthetic dataset generated:")
    for k, v in paths.items():
        print(f"- {k}: {v}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
