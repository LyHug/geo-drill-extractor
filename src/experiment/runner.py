"""
实验执行器 - 批量运行多模型多文档实验
"""

import os
import json
import csv
import logging
import traceback
from collections import Counter
from datetime import datetime
from pathlib import Path
from threading import Lock
from typing import Dict, List, Optional, Any, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
import time

from core import (
    LLMModel,
    ProcessResult,
    ConfigLoader,
    get_config_loader,
    ExperimentException
)
from extraction import ExtractionPipeline
from evaluation import (
    SixMetricsProcessor,
    GroundTruthLoader,
    get_tokenizer_manager
)

logger = logging.getLogger(__name__)


class ExperimentRunner:
    """
    实验执行器
    
    负责批量运行多模型多文档的提取实验，支持：
    - 多模型并行处理
    - 多轮重复实验
    - 实验结果聚合和导出
    - 错误处理和恢复
    """
    
    def __init__(
        self,
        models: Optional[List[LLMModel]] = None,
        config_loader: Optional[ConfigLoader] = None,
        stream_mode: bool = False
    ):
        """
        初始化实验执行器
        
        Args:
            models: 要测试的模型列表，如果为None则使用配置中的默认模型
            config_loader: 配置加载器
            stream_mode: 是否启用流式输出显示
        """
        self.config_loader = config_loader or get_config_loader()
        self.stream_mode = stream_mode
        
        # 设置模型列表
        if models:
            self.models = models
        else:
            # 使用配置中的模型列表
            default_models = self.config_loader.get('experiment.default_models', [])
            self.models = [LLMModel(model) for model in default_models] if default_models else [LLMModel.DEEPSEEK_R1_DISTILL_QWEN_32B_ALIYUN]
        
        # 获取数据路径配置
        data_paths = self.config_loader.get_data_paths()
        
        # 加载配置
        self.documents_dir = Path(data_paths['documents_dir'])
        output_dir_value = self.config_loader.get('experiment.output_dir', './experiment_results')
        output_dir_path = Path(str(output_dir_value))
        if not output_dir_path.is_absolute():
            try:
                config_path = Path(self.config_loader.config_path)
                if config_path.exists():
                    project_root = config_path.resolve().parent.parent  # configs/config.yaml -> 项目根目录
                    output_dir_path = project_root / output_dir_path
            except Exception:
                pass
        self.output_dir = output_dir_path
        self.max_workers = self.config_loader.get('processing.parallel.max_workers', 3)
        
        # 初始化组件 - 使用配置的路径
        self.ground_truth_loader = GroundTruthLoader(
            annotations_file=data_paths['ground_truth_file']
        )
        self.metrics_processor = SixMetricsProcessor(self.ground_truth_loader)
        self.tokenizer_manager = get_tokenizer_manager()
        self._stats_lock = Lock()
        
        # 实验状态
        self.current_experiment_dir = None
        self.experiment_stats = {
            'start_time': None,
            'end_time': None,
            'total_documents': 0,
            'total_models': 0,
            'total_repetitions': 0,
            'successful_runs': 0,
            'failed_runs': 0,
            'errors': []
        }

    @staticmethod
    def _json_default(obj: Any):
        """JSON序列化兜底（处理numpy标量等非原生类型）。"""
        if hasattr(obj, 'item') and callable(getattr(obj, 'item', None)):
            try:
                return obj.item()
            except Exception:
                pass
        if isinstance(obj, datetime):
            return obj.isoformat()
        return str(obj)
    
    def run_experiment(
        self,
        repetitions: int = 5,
        test_documents: Optional[int] = None,
        document_filter: Optional[str] = None,
        save_results: bool = True
    ) -> Dict[str, Any]:
        """
        运行完整实验
        
        Args:
            repetitions: 每个模型每个文档的重复次数
            test_documents: 限制测试文档数量（用于快速测试）
            document_filter: 文档名称过滤器（包含指定字符串的文档）
            save_results: 是否保存结果到文件
        
        Returns:
            实验结果字典
        """
        logger.info(f"开始实验: {len(self.models)} 个模型, {repetitions} 轮重复")
        
        # 初始化实验
        self._initialize_experiment()
        
        try:
            # 获取文档列表
            documents = self._get_document_list(test_documents, document_filter)
            logger.info(f"找到 {len(documents)} 个待处理文档")
            
            # 更新统计信息
            self.experiment_stats.update({
                'start_time': datetime.now().isoformat(),
                'total_documents': len(documents),
                'total_models': len(self.models),
                'total_repetitions': repetitions
            })
            
            # 批量处理文档
            all_results = {}
            documents_content = {}

            # 预加载文档内容（用于token计算；避免并发读取同一文件）
            for doc_path in documents:
                documents_content[doc_path.name] = self._load_document_content(doc_path)
            
            for model in self.models:
                logger.info(f"开始处理模型: {model.value}")
                
                model_results = []

                if self.max_workers and self.max_workers > 1 and len(documents) > 1:
                    # 文档级并行（同一模型下并行处理多个文档）。注意：并行会触发更多并发LLM请求，
                    # 如出现限流/超时可在 configs/config.yaml 下调 processing.parallel.max_workers。
                    with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                        future_to_doc = {
                            executor.submit(self._run_document_experiments, model, doc_path, repetitions): doc_path
                            for doc_path in documents
                        }

                        for future in as_completed(future_to_doc):
                            doc_path = future_to_doc[future]
                            try:
                                doc_results = future.result()
                                model_results.extend(doc_results)
                            except Exception as e:
                                error_msg = f"处理文档 {doc_path.name} 失败: {str(e)}"
                                logger.error(error_msg)
                                with self._stats_lock:
                                    self.experiment_stats['errors'].append(error_msg)
                                    self.experiment_stats['failed_runs'] += repetitions
                else:
                    # 串行处理（更稳健，适合限流严格或超时较多的模型）
                    for doc_path in documents:
                        try:
                            doc_results = self._run_document_experiments(
                                model, doc_path, repetitions
                            )
                            model_results.extend(doc_results)
                        except Exception as e:
                            error_msg = f"处理文档 {doc_path.name} 失败: {str(e)}"
                            logger.error(error_msg)
                            with self._stats_lock:
                                self.experiment_stats['errors'].append(error_msg)
                                self.experiment_stats['failed_runs'] += repetitions
                
                all_results[model.value] = model_results
            
            # 计算指标
            logger.info("计算评估指标...")
            metrics_results = self.metrics_processor.batch_process_results(
                all_results, documents_content
            )
            
            # 组装最终结果
            experiment_results = {
                'metadata': {
                    'experiment_time': self.experiment_stats['start_time'],
                    'models': [model.value for model in self.models],
                    'total_documents': len(documents),
                    'repetitions': repetitions,
                    'config_snapshot': self._get_config_snapshot()
                },
                'raw_results': all_results,
                'metrics': metrics_results,
                'statistics': self.experiment_stats
            }
            
            # 保存结果
            if save_results and self.current_experiment_dir:
                self._save_experiment_results(experiment_results)
            
            # 更新完成时间
            self.experiment_stats['end_time'] = datetime.now().isoformat()
            
            logger.info("实验完成!")
            self._log_experiment_summary()
            
            return experiment_results
            
        except Exception as e:
            logger.error(f"实验执行失败: {str(e)}")
            logger.error(traceback.format_exc())
            self.experiment_stats['end_time'] = datetime.now().isoformat()
            raise ExperimentException(f"Experiment failed: {str(e)}")
    
    def _initialize_experiment(self):
        """初始化实验目录和日志"""
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        self.current_experiment_dir = self.output_dir / timestamp
        self.current_experiment_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"实验输出目录: {self.current_experiment_dir}")
        
        # 重置统计信息
        self.experiment_stats = {
            'start_time': None,
            'end_time': None,
            'total_documents': 0,
            'total_models': 0,
            'total_repetitions': 0,
            'successful_runs': 0,
            'failed_runs': 0,
            'errors': []
        }
    
    def _get_document_list(
        self,
        limit: Optional[int] = None,
        name_filter: Optional[str] = None
    ) -> List[Path]:
        """获取待处理文档列表"""
        documents = []
        
        # 搜索docx文件
        for doc_path in self.documents_dir.glob("*.docx"):
            if name_filter and name_filter not in doc_path.name:
                continue
            documents.append(doc_path)
        
        # 排序确保结果一致性
        documents.sort(key=lambda x: x.name)
        
        # 应用数量限制
        if limit and limit > 0:
            documents = documents[:limit]
        
        return documents
    
    def _load_document_content(self, doc_path: Path) -> str:
        """加载文档内容用于token计算"""
        try:
            from docx import Document
            doc = Document(doc_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        content.append(' | '.join(row_data))
            
            return '\n'.join(content)
            
        except Exception as e:
            logger.warning(f"加载文档内容失败 {doc_path.name}: {str(e)}")
            return ""
    
    def _run_document_experiments(
        self,
        model: LLMModel,
        doc_path: Path,
        repetitions: int
    ) -> List[ProcessResult]:
        """运行单个文档的多轮实验"""
        results = []
        
        for rep in range(repetitions):
            try:
                logger.info(f"处理 {doc_path.name} - 模型 {model.value} - 第 {rep + 1}/{repetitions} 轮")
                
                # 创建提取管道
                pipeline = ExtractionPipeline(model=model, stream_mode=self.stream_mode)
                
                # 处理文档
                start_time = time.time()
                result = pipeline.process_document(str(doc_path))
                processing_time = time.time() - start_time
                
                # 添加元数据
                if result.metadata is None:
                    result.metadata = {}
                
                result.metadata.update({
                    'repetition_round': rep + 1,
                    'model_name': model.value,
                    'experiment_timestamp': datetime.now().isoformat()
                })
                
                result.processing_time = processing_time
                results.append(result)
                
                with self._stats_lock:
                    self.experiment_stats['successful_runs'] += 1
                logger.info(f"成功完成 {doc_path.name} 第 {rep + 1} 轮，用时 {processing_time:.2f}s")
                
            except Exception as e:
                error_msg = f"处理失败 {doc_path.name} 第 {rep + 1} 轮: {str(e)}"
                logger.error(error_msg)
                
                with self._stats_lock:
                    self.experiment_stats['failed_runs'] += 1
                    self.experiment_stats['errors'].append(error_msg)
                
                # 创建错误结果
                error_result = ProcessResult(
                    document_name=doc_path.name,
                    drill_holes=[],
                    coordinates={},
                    processing_time=0,
                    errors=[{
                        'type': 'ProcessingError',
                        'message': str(e),
                        'traceback': traceback.format_exc()
                    }],
                    metadata={
                        'repetition_round': rep + 1,
                        'model_name': model.value,
                        'error_occurred': True
                    }
                )
                results.append(error_result)
        
        return results
    
    def _get_config_snapshot(self) -> Dict[str, Any]:
        """获取当前配置快照"""
        return {
            'models': [model.value for model in self.models],
            'documents_dir': str(self.documents_dir),
            'output_dir': str(self.output_dir),
            'max_workers': self.max_workers,
            'llm_config': {
                'timeout': self.config_loader.get('llm.timeout', 60),
                'max_retries': self.config_loader.get('llm.max_retries', 3),
                'temperature': self.config_loader.get('llm.temperature', 0.1)
            }
        }
    
    def _save_experiment_results(self, results: Dict[str, Any]):
        """保存实验结果到文件"""
        try:
            from .exporter import JSONExporter, FieldMapper

            export_time = datetime.now().isoformat()
            raw_results_by_model: Dict[str, List[ProcessResult]] = results.get('raw_results', {})
            metrics_by_model = results.get('metrics', {})
            metadata = results.get('metadata', {})
            statistics = results.get('statistics', {})

            # 1) 导出原始提取结果（包含metadata，便于后续失败模式统计）
            json_exporter = JSONExporter(self.config_loader)
            raw_results_export = {
                'export_info': {
                    'export_time': export_time,
                    'total_models': len(raw_results_by_model),
                    'total_runs': sum(len(v) for v in raw_results_by_model.values()),
                },
                'models': {}
            }
            for model_name, model_results in raw_results_by_model.items():
                raw_results_export['models'][model_name] = json_exporter._prepare_results_json_data(model_results)

            raw_results_file = self.current_experiment_dir / "raw_results.json"
            with open(raw_results_file, 'w', encoding='utf-8') as f:
                json.dump(raw_results_export, f, ensure_ascii=False, indent=2, default=self._json_default)

            # 2) 导出指标结果（用于统计分析/绘图）
            field_mapper = FieldMapper(self.config_loader)
            metrics_export = {
                'export_info': {
                    'export_time': export_time,
                    'total_models': len(metrics_by_model),
                    'total_documents': metadata.get('total_documents'),
                    'repetitions': metadata.get('repetitions'),
                },
                'models': {}
            }
            metrics_rows: List[Dict[str, Any]] = []
            for model_name, model_metrics in metrics_by_model.items():
                mapped = [field_mapper.map_metrics_to_dict(m) for m in model_metrics]
                metrics_export['models'][model_name] = mapped
                metrics_rows.extend(mapped)

            metrics_file = self.current_experiment_dir / "metrics_results.json"
            with open(metrics_file, 'w', encoding='utf-8') as f:
                json.dump(metrics_export, f, ensure_ascii=False, indent=2, default=self._json_default)

            metrics_csv_file = None
            if metrics_rows:
                metrics_csv_file = self.current_experiment_dir / "metrics_results.csv"
                fieldnames = sorted({k for row in metrics_rows for k in row.keys()})
                with open(metrics_csv_file, 'w', newline='', encoding='utf-8-sig') as csvfile:
                    writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                    writer.writeheader()
                    writer.writerows(metrics_rows)

            # 3) 失败模式汇总（区分协议/解析失败 vs 语义/几何失败）
            failure_summary, failure_rows = self._summarize_failure_modes(raw_results_by_model)
            failure_summary_file = self.current_experiment_dir / "failure_modes_summary.json"
            with open(failure_summary_file, 'w', encoding='utf-8') as f:
                json.dump(failure_summary, f, ensure_ascii=False, indent=2, default=self._json_default)

            failure_csv_file = None
            if failure_rows:
                failure_csv_file = self.current_experiment_dir / "failure_modes_breakdown.csv"
                with open(failure_csv_file, 'w', newline='', encoding='utf-8-sig') as csvfile:
                    writer = csv.DictWriter(
                        csvfile,
                        fieldnames=['model_name', 'stage', 'reason', 'count', 'denominator', 'proportion'],
                    )
                    writer.writeheader()
                    writer.writerows(failure_rows)

            # 4) 保存实验总览（包含指向导出文件的引用）
            results_file = self.current_experiment_dir / "experiment_results.json"
            export_payload = {
                'metadata': metadata,
                'statistics': statistics,
                'raw_results': raw_results_export,
                'metrics': metrics_export,
                'failure_modes': failure_summary,
                'exported_files': {
                    'raw_results_json': str(raw_results_file.name),
                    'metrics_results_json': str(metrics_file.name),
                    'metrics_results_csv': str(metrics_csv_file.name) if metrics_csv_file else None,
                    'failure_modes_summary_json': str(failure_summary_file.name),
                    'failure_modes_breakdown_csv': str(failure_csv_file.name) if failure_csv_file else None,
                }
            }
            with open(results_file, 'w', encoding='utf-8') as f:
                json.dump(export_payload, f, ensure_ascii=False, indent=2, default=self._json_default)

            # 5) 生成便于查看的摘要文本
            summary_file = self.current_experiment_dir / "processing_summary.txt"
            successful_runs = int(statistics.get('successful_runs') or 0)
            failed_runs = int(statistics.get('failed_runs') or 0)
            total_runs = successful_runs + failed_runs
            success_rate = (successful_runs / total_runs * 100) if total_runs else 0.0

            with open(summary_file, 'w', encoding='utf-8') as f:
                f.write("Experiment Summary\n")
                f.write("=" * 60 + "\n")
                f.write(f"export_time: {export_time}\n")
                f.write(f"models: {len(metadata.get('models', []))}\n")
                f.write(f"documents: {metadata.get('total_documents')}\n")
                f.write(f"repetitions: {metadata.get('repetitions')}\n")
                f.write(f"successful_runs: {successful_runs}\n")
                f.write(f"failed_runs: {failed_runs}\n")
                f.write(f"success_rate_percent: {success_rate:.2f}\n")
                f.write("\nFiles\n")
                f.write("-" * 60 + "\n")
                f.write(f"experiment_results.json: {results_file.name}\n")
                f.write(f"raw_results.json: {raw_results_file.name}\n")
                f.write(f"metrics_results.json: {metrics_file.name}\n")
                if metrics_csv_file:
                    f.write(f"metrics_results.csv: {metrics_csv_file.name}\n")
                f.write(f"failure_modes_summary.json: {failure_summary_file.name}\n")
                if failure_csv_file:
                    f.write(f"failure_modes_breakdown.csv: {failure_csv_file.name}\n")

            logger.info(f"实验结果已保存到: {self.current_experiment_dir}")
            
        except Exception as e:
            logger.error(f"保存实验结果失败: {str(e)}")
            logger.error(traceback.format_exc())

    def _summarize_failure_modes(
        self,
        results_by_model: Dict[str, List[ProcessResult]],
    ) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
        """
        汇总失败模式（基于每次运行metadata中的失败计数）。

        Returns:
            (summary_json, breakdown_rows_for_csv)
        """
        export_time = datetime.now().isoformat()
        summary: Dict[str, Any] = {
            'export_time': export_time,
            'models': {}
        }
        breakdown_rows: List[Dict[str, Any]] = []

        for model_name, results in results_by_model.items():
            total_runs = len(results)

            unique_location_total = 0
            location_success_total = 0
            location_failure_counts: Counter = Counter()
            start_coordinate_failure_counts: Counter = Counter()

            for result in results:
                meta = result.metadata or {}

                unique_location_total += int(meta.get('unique_location_descriptions_count') or 0)
                location_success_total += int(meta.get('location_analysis_success_count') or 0)

                location_failures = meta.get('location_analysis_failure_counts') or {}
                if isinstance(location_failures, dict):
                    location_failure_counts.update({
                        str(k): int(v) for k, v in location_failures.items()
                    })

                start_failures = meta.get('start_coordinate_failure_counts') or {}
                if isinstance(start_failures, dict):
                    start_coordinate_failure_counts.update({
                        str(k): int(v) for k, v in start_failures.items()
                    })

            location_failure_total = int(sum(location_failure_counts.values()))
            location_attempt_total = int(location_success_total + location_failure_total)

            summary['models'][model_name] = {
                'total_runs': total_runs,
                'unique_location_descriptions_total': unique_location_total,
                'location_analysis_success_total': location_success_total,
                'location_analysis_failure_total': location_failure_total,
                'location_analysis_attempt_total': location_attempt_total,
                'location_analysis_failure_counts': dict(location_failure_counts),
                'start_coordinate_failure_counts': dict(start_coordinate_failure_counts),
            }

            for reason, count in location_failure_counts.items():
                denom = location_attempt_total or 0
                breakdown_rows.append({
                    'model_name': model_name,
                    'stage': 'location_analysis',
                    'reason': reason,
                    'count': int(count),
                    'denominator': int(denom),
                    'proportion': (float(count) / denom) if denom else None,
                })

            for reason, count in start_coordinate_failure_counts.items():
                denom = location_success_total or 0
                breakdown_rows.append({
                    'model_name': model_name,
                    'stage': 'start_coordinate',
                    'reason': reason,
                    'count': int(count),
                    'denominator': int(denom),
                    'proportion': (float(count) / denom) if denom else None,
                })

        return summary, breakdown_rows
    
    def _log_experiment_summary(self):
        """记录实验摘要"""
        stats = self.experiment_stats
        
        total_runs = stats['successful_runs'] + stats['failed_runs']
        success_rate = (stats['successful_runs'] / total_runs * 100) if total_runs > 0 else 0
        
        if stats['start_time'] and stats['end_time']:
            start = datetime.fromisoformat(stats['start_time'])
            end = datetime.fromisoformat(stats['end_time'])
            duration = (end - start).total_seconds()
        else:
            duration = 0
        
        logger.info("="*60)
        logger.info("实验摘要:")
        logger.info(f"  总文档数: {stats['total_documents']}")
        logger.info(f"  总模型数: {stats['total_models']}")
        logger.info(f"  重复轮次: {stats['total_repetitions']}")
        logger.info(f"  成功运行: {stats['successful_runs']}")
        logger.info(f"  失败运行: {stats['failed_runs']}")
        logger.info(f"  成功率: {success_rate:.1f}%")
        logger.info(f"  总耗时: {duration:.1f}s")
        logger.info(f"  错误数量: {len(stats['errors'])}")
        logger.info("="*60)


def run_quick_experiment(
    models: Optional[List[str]] = None,
    test_documents: int = 5,
    repetitions: int = 2
) -> Dict[str, Any]:
    """
    运行快速测试实验
    
    Args:
        models: 模型名称列表
        test_documents: 测试文档数量
        repetitions: 重复次数
    
    Returns:
        实验结果
    """
    logger.info("运行快速实验...")
    
    # 转换模型名称为枚举
    model_enums = []
    if models:
        for model_name in models:
            try:
                model_enums.append(LLMModel(model_name))
            except ValueError:
                logger.warning(f"未知模型: {model_name}")
    else:
        model_enums = [LLMModel.DEEPSEEK_R1_DISTILL_QWEN_32B_ALIYUN]
    
    # 创建实验执行器并运行
    runner = ExperimentRunner(models=model_enums)
    return runner.run_experiment(
        repetitions=repetitions,
        test_documents=test_documents,
        save_results=True
    )


def run_full_experiment(
    models: Optional[List[str]] = None,
    repetitions: int = 5
) -> Dict[str, Any]:
    """
    运行完整实验
    
    Args:
        models: 模型名称列表
        repetitions: 重复次数
    
    Returns:
        实验结果
    """
    logger.info("运行完整实验...")
    
    # 转换模型名称为枚举
    model_enums = []
    if models:
        for model_name in models:
            try:
                model_enums.append(LLMModel(model_name))
            except ValueError:
                logger.warning(f"未知模型: {model_name}")
    else:
        # 使用默认模型集合
        model_enums = [
            LLMModel.DEEPSEEK_R1_DISTILL_QWEN_32B_ALIYUN,
            LLMModel.DEEPSEEK_V3_OFFICIAL,
            LLMModel.QWEN3_32B,
            LLMModel.GPT_35_TURBO_OPENROUTER
        ]
    
    # 创建实验执行器并运行
    runner = ExperimentRunner(models=model_enums)
    return runner.run_experiment(
        repetitions=repetitions,
        save_results=True
    )
