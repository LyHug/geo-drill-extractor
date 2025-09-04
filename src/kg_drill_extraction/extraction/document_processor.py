"""
文档处理器模块 - 处理Word文档并转换为结构化文本
"""

import re
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..core.exceptions import (
    DocumentException,
    DocumentNotFoundException,
    DocumentReadException,
    DocumentParseException,
    UnsupportedDocumentFormatException
)

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    文档处理器
    
    负责将各种格式的文档转换为结构化文本，目前主要支持Word文档
    """
    
    def __init__(self):
        """初始化文档处理器"""
        self.supported_formats = ['.docx', '.doc']
    
    def process_document(self, file_path: str) -> str:
        """
        处理文档文件
        
        Args:
            file_path: 文档文件路径
        
        Returns:
            处理后的文本内容
        
        Raises:
            DocumentException: 文档处理异常
        """
        path = Path(file_path)
        
        # 检查文件是否存在
        if not path.exists():
            raise DocumentNotFoundException(
                f"Document not found: {file_path}",
                details={'path': str(path.absolute())}
            )
        
        # 检查文件格式
        if path.suffix.lower() not in self.supported_formats:
            raise UnsupportedDocumentFormatException(
                f"Unsupported format: {path.suffix}",
                details={
                    'file': file_path,
                    'supported_formats': self.supported_formats
                }
            )
        
        # 处理Word文档
        if path.suffix.lower() == '.docx':
            return self.convert_docx_to_markdown(path)
        else:
            raise UnsupportedDocumentFormatException(
                f"Format {path.suffix} not implemented yet"
            )
    
    def convert_docx_to_markdown(self, file_path: Path) -> str:
        """
        将Word文档转换为Markdown格式
        
        Args:
            file_path: Word文档路径
        
        Returns:
            Markdown格式的文本
        
        Raises:
            DocumentReadException: 读取文档失败
            DocumentParseException: 解析文档失败
        """
        try:
            doc = Document(file_path)
            markdown_content = []
            
            # 遍历文档元素
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # 处理段落
                    paragraph = Paragraph(element, doc)
                    text = paragraph.text.strip()
                    if text:
                        # 识别标题
                        if paragraph.style and paragraph.style.name:
                            if paragraph.style.name.startswith('Heading'):
                                level = self._get_heading_level(paragraph.style.name)
                                markdown_content.append('#' * level + ' ' + text)
                            else:
                                markdown_content.append(text)
                        else:
                            markdown_content.append(text)
                        markdown_content.append('')  # 添加空行
                
                elif isinstance(element, CT_Tbl):
                    # 处理表格
                    table = Table(element, doc)
                    markdown_table = self._process_table(table)
                    markdown_content.extend(markdown_table)
            
            result = '\n'.join(markdown_content)
            
            # 清理多余的空行
            result = re.sub(r'\n{3,}', '\n\n', result)
            
            logger.info(f"成功转换文档 {file_path.name}, 内容长度: {len(result)}")
            return result
            
        except Exception as e:
            if isinstance(e, DocumentException):
                raise
            elif "Document" in str(e) or "docx" in str(e):
                raise DocumentReadException(
                    f"Failed to read document: {str(e)}",
                    details={'file': str(file_path)}
                )
            else:
                raise DocumentParseException(
                    f"Failed to parse document: {str(e)}",
                    details={'file': str(file_path)}
                )
    
    def _get_heading_level(self, style_name: str) -> int:
        """
        获取标题级别
        
        Args:
            style_name: 样式名称
        
        Returns:
            标题级别（1-6）
        """
        # 从样式名称中提取数字
        match = re.search(r'(\d+)', style_name)
        if match:
            level = int(match.group(1))
            return min(level, 6)  # 限制最大级别为6
        return 1
    
    def _process_table(self, table: Table) -> List[str]:
        """
        处理表格并转换为Markdown格式
        
        Args:
            table: Word表格对象
        
        Returns:
            Markdown格式的表格行列表
        """
        if not table.rows:
            return []
        
        markdown_table = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                # 合并单元格中的所有段落文本
                cell_text = ' '.join(
                    p.text.strip() for p in cell.paragraphs if p.text.strip()
                )
                # 清理换行符和多余空格
                cell_text = re.sub(r'\s+', ' ', cell_text)
                row_data.append(cell_text)
            
            # 创建Markdown表格行
            markdown_row = '| ' + ' | '.join(row_data) + ' |'
            markdown_table.append(markdown_row)
            
            # 在第一行后添加分隔符
            if row_idx == 0:
                separator = '|' + '---|' * len(row_data)
                markdown_table.append(separator)
        
        markdown_table.append('')  # 添加空行
        return markdown_table
    
    def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """
        提取文档中的所有表格
        
        Args:
            file_path: 文档文件路径
        
        Returns:
            表格数据列表
        """
        path = Path(file_path)
        
        if not path.exists():
            raise DocumentNotFoundException(f"Document not found: {file_path}")
        
        if path.suffix.lower() != '.docx':
            raise UnsupportedDocumentFormatException(
                f"Table extraction only supports .docx format"
            )
        
        try:
            doc = Document(path)
            tables_data = []
            
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'index': table_idx,
                    'rows': [],
                    'headers': None
                }
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = ' '.join(
                            p.text.strip() for p in cell.paragraphs if p.text.strip()
                        )
                        row_data.append(cell_text)
                    
                    if row_idx == 0:
                        table_data['headers'] = row_data
                    
                    table_data['rows'].append(row_data)
                
                tables_data.append(table_data)
            
            return tables_data
            
        except Exception as e:
            raise DocumentParseException(
                f"Failed to extract tables: {str(e)}",
                details={'file': str(path)}
            )
    
    def get_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        获取文档元数据
        
        Args:
            file_path: 文档文件路径
        
        Returns:
            文档元数据字典
        """
        path = Path(file_path)
        
        if not path.exists():
            raise DocumentNotFoundException(f"Document not found: {file_path}")
        
        metadata = {
            'file_name': path.name,
            'file_size': path.stat().st_size,
            'file_path': str(path.absolute()),
            'format': path.suffix.lower()
        }
        
        if path.suffix.lower() == '.docx':
            try:
                doc = Document(path)
                
                # 获取文档属性
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'created': core_props.created.isoformat() if core_props.created else None,
                    'modified': core_props.modified.isoformat() if core_props.modified else None,
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                })
            except Exception as e:
                logger.warning(f"Failed to extract document properties: {str(e)}")
        
        return metadata