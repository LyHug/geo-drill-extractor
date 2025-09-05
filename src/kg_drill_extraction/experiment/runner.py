"""
实验执行器 - 批量运行多模型多文档实验
"""

import os
import json
import logging
import traceback
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any
from concurrent.futures import ThreadPoolExecutor, as_completed
import time

from ..core import (
    LLMModel,
    ProcessResult,
    ConfigLoader,
    get_config_loader,
    ExperimentException
)
from ..extraction import ExtractionPipeline
from ..evaluation import (
    SixMetricsProcessor,
    GroundTruthLoader,
    get_tokenizer_manager
)

logger = logging.getLogger(__name__)


class ExperimentRunner:
    """
    实验执行器
    
    负责批量运行多模型多文档的提取实验，支持：
    - 多模型并行处理
    - 多轮重复实验
    - 实验结果聚合和导出
    - 错误处理和恢复
    """
    
    def __init__(
        self,
        models: Optional[List[LLMModel]] = None,
        config_loader: Optional[ConfigLoader] = None,
        stream_mode: bool = False
    ):
        """
        初始化实验执行器
        
        Args:
            models: 要测试的模型列表，如果为None则使用配置中的默认模型
            config_loader: 配置加载器
            stream_mode: 是否启用流式输出显示
        """
        self.config_loader = config_loader or get_config_loader()
        self.stream_mode = stream_mode
        
        # 设置模型列表
        if models:
            self.models = models
        else:
            # 使用配置中的模型列表
            default_models = self.config_loader.get('experiment.default_models', [])
            self.models = [LLMModel(model) for model in default_models] if default_models else [LLMModel.DEEPSEEK_R1_DISTILL_QWEN_32B_ALIYUN]
        
        # 获取数据路径配置
        data_paths = self.config_loader.get_data_paths()
        
        # 加载配置
        self.documents_dir = Path(data_paths['documents_dir'])
        self.output_dir = Path(self.config_loader.get('experiment.output_dir', './experiment_results'))
        self.max_workers = self.config_loader.get('processing.parallel.max_workers', 3)
        
        # 初始化组件 - 使用配置的路径
        self.ground_truth_loader = GroundTruthLoader(
            annotations_file=data_paths['ground_truth_file']
        )
        self.metrics_processor = SixMetricsProcessor(self.ground_truth_loader)
        self.tokenizer_manager = get_tokenizer_manager()
        
        # 实验状态
        self.current_experiment_dir = None
        self.experiment_stats = {
            'start_time': None,
            'end_time': None,
            'total_documents': 0,
            'total_models': 0,
            'total_repetitions': 0,
            'successful_runs': 0,
            'failed_runs': 0,
            'errors': []
        }
    
    def run_experiment(
        self,
        repetitions: int = 5,
        test_documents: Optional[int] = None,
        document_filter: Optional[str] = None,
        save_results: bool = True
    ) -> Dict[str, Any]:
        """
        运行完整实验
        
        Args:
            repetitions: 每个模型每个文档的重复次数
            test_documents: 限制测试文档数量（用于快速测试）
            document_filter: 文档名称过滤器（包含指定字符串的文档）
            save_results: 是否保存结果到文件
        
        Returns:
            实验结果字典
        """
        logger.info(f"开始实验: {len(self.models)} 个模型, {repetitions} 轮重复")
        
        # 初始化实验
        self._initialize_experiment()
        
        try:
            # 获取文档列表
            documents = self._get_document_list(test_documents, document_filter)
            logger.info(f"找到 {len(documents)} 个待处理文档")
            
            # 更新统计信息
            self.experiment_stats.update({
                'start_time': datetime.now().isoformat(),
                'total_documents': len(documents),
                'total_models': len(self.models),
                'total_repetitions': repetitions
            })
            
            # 批量处理文档
            all_results = {}
            documents_content = {}
            
            for model in self.models:
                logger.info(f"开始处理模型: {model.value}")
                
                model_results = []
                for doc_path in documents:
                    try:
                        # 加载文档内容（用于token计算）
                        if doc_path.name not in documents_content:
                            documents_content[doc_path.name] = self._load_document_content(doc_path)
                        
                        # 运行多轮实验
                        doc_results = self._run_document_experiments(
                            model, doc_path, repetitions
                        )
                        model_results.extend(doc_results)
                        
                    except Exception as e:
                        error_msg = f"处理文档 {doc_path.name} 失败: {str(e)}"
                        logger.error(error_msg)
                        self.experiment_stats['errors'].append(error_msg)
                        self.experiment_stats['failed_runs'] += repetitions
                
                all_results[model.value] = model_results
            
            # 计算指标
            logger.info("计算评估指标...")
            metrics_results = self.metrics_processor.batch_process_results(
                all_results, documents_content
            )
            
            # 组装最终结果
            experiment_results = {
                'metadata': {
                    'experiment_time': self.experiment_stats['start_time'],
                    'models': [model.value for model in self.models],
                    'total_documents': len(documents),
                    'repetitions': repetitions,
                    'config_snapshot': self._get_config_snapshot()
                },
                'raw_results': all_results,
                'metrics': metrics_results,
                'statistics': self.experiment_stats
            }
            
            # 保存结果
            if save_results and self.current_experiment_dir:
                self._save_experiment_results(experiment_results)
            
            # 更新完成时间
            self.experiment_stats['end_time'] = datetime.now().isoformat()
            
            logger.info("实验完成!")
            self._log_experiment_summary()
            
            return experiment_results
            
        except Exception as e:
            logger.error(f"实验执行失败: {str(e)}")
            logger.error(traceback.format_exc())
            self.experiment_stats['end_time'] = datetime.now().isoformat()
            raise ExperimentException(f"Experiment failed: {str(e)}")
    
    def _initialize_experiment(self):
        """初始化实验目录和日志"""
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        self.current_experiment_dir = self.output_dir / timestamp
        self.current_experiment_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"实验输出目录: {self.current_experiment_dir}")
        
        # 重置统计信息
        self.experiment_stats = {
            'start_time': None,
            'end_time': None,
            'total_documents': 0,
            'total_models': 0,
            'total_repetitions': 0,
            'successful_runs': 0,
            'failed_runs': 0,
            'errors': []
        }
    
    def _get_document_list(
        self,
        limit: Optional[int] = None,
        name_filter: Optional[str] = None
    ) -> List[Path]:
        """获取待处理文档列表"""
        documents = []
        
        # 搜索docx文件
        for doc_path in self.documents_dir.glob("*.docx"):
            if name_filter and name_filter not in doc_path.name:
                continue
            documents.append(doc_path)
        
        # 排序确保结果一致性
        documents.sort(key=lambda x: x.name)
        
        # 应用数量限制
        if limit and limit > 0:
            documents = documents[:limit]
        
        return documents
    
    def _load_document_content(self, doc_path: Path) -> str:
        """加载文档内容用于token计算"""
        try:
            from docx import Document
            doc = Document(doc_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        content.append(' | '.join(row_data))
            
            return '\n'.join(content)
            
        except Exception as e:
            logger.warning(f"加载文档内容失败 {doc_path.name}: {str(e)}")
            return ""
    
    def _run_document_experiments(
        self,
        model: LLMModel,
        doc_path: Path,
        repetitions: int
    ) -> List[ProcessResult]:
        """运行单个文档的多轮实验"""
        results = []
        
        for rep in range(repetitions):
            try:
                logger.info(f"处理 {doc_path.name} - 模型 {model.value} - 第 {rep + 1}/{repetitions} 轮")
                
                # 创建提取管道
                pipeline = ExtractionPipeline(model=model, stream_mode=self.stream_mode)
                
                # 处理文档
                start_time = time.time()
                result = pipeline.process_document(str(doc_path))
                processing_time = time.time() - start_time
                
                # 添加元数据
                if result.metadata is None:
                    result.metadata = {}
                
                result.metadata.update({
                    'repetition_round': rep + 1,
                    'model_name': model.value,
                    'experiment_timestamp': datetime.now().isoformat()
                })
                
                result.processing_time = processing_time
                results.append(result)
                
                self.experiment_stats['successful_runs'] += 1
                logger.info(f"成功完成 {doc_path.name} 第 {rep + 1} 轮，用时 {processing_time:.2f}s")
                
            except Exception as e:
                error_msg = f"处理失败 {doc_path.name} 第 {rep + 1} 轮: {str(e)}"
                logger.error(error_msg)
                
                self.experiment_stats['failed_runs'] += 1
                self.experiment_stats['errors'].append(error_msg)
                
                # 创建错误结果
                error_result = ProcessResult(
                    document_name=doc_path.name,
                    drill_holes=[],
                    coordinates=[],
                    processing_time=0,
                    errors=[str(e)],
                    metadata={
                        'repetition_round': rep + 1,
                        'model_name': model.value,
                        'error_occurred': True
                    }
                )
                results.append(error_result)
        
        return results
    
    def _get_config_snapshot(self) -> Dict[str, Any]:
        """获取当前配置快照"""
        return {
            'models': [model.value for model in self.models],
            'documents_dir': str(self.documents_dir),
            'output_dir': str(self.output_dir),
            'max_workers': self.max_workers,
            'llm_config': {
                'timeout': self.config_loader.get('llm.timeout', 60),
                'max_retries': self.config_loader.get('llm.max_retries', 3),
                'temperature': self.config_loader.get('llm.temperature', 0.1)
            }
        }
    
    def _save_experiment_results(self, results: Dict[str, Any]):
        """保存实验结果到文件"""
        try:
            # 保存完整结果为JSON
            results_file = self.current_experiment_dir / "experiment_results.json"
            with open(results_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, ensure_ascii=False, indent=2, default=str)
            
            # 保存处理结果摘要
            raw_results_file = self.current_experiment_dir / "raw_results.json"
            with open(raw_results_file, 'w', encoding='utf-8') as f:
                json.dump(results['raw_results'], f, ensure_ascii=False, indent=2, default=str)
            
            # 保存指标结果
            metrics_file = self.current_experiment_dir / "metrics_results.json"
            with open(metrics_file, 'w', encoding='utf-8') as f:
                json.dump(results['metrics'], f, ensure_ascii=False, indent=2, default=str)
            
            logger.info(f"实验结果已保存到: {self.current_experiment_dir}")
            
        except Exception as e:
            logger.error(f"保存实验结果失败: {str(e)}")
    
    def _log_experiment_summary(self):
        """记录实验摘要"""
        stats = self.experiment_stats
        
        total_runs = stats['successful_runs'] + stats['failed_runs']
        success_rate = (stats['successful_runs'] / total_runs * 100) if total_runs > 0 else 0
        
        if stats['start_time'] and stats['end_time']:
            start = datetime.fromisoformat(stats['start_time'])
            end = datetime.fromisoformat(stats['end_time'])
            duration = (end - start).total_seconds()
        else:
            duration = 0
        
        logger.info("="*60)
        logger.info("实验摘要:")
        logger.info(f"  总文档数: {stats['total_documents']}")
        logger.info(f"  总模型数: {stats['total_models']}")
        logger.info(f"  重复轮次: {stats['total_repetitions']}")
        logger.info(f"  成功运行: {stats['successful_runs']}")
        logger.info(f"  失败运行: {stats['failed_runs']}")
        logger.info(f"  成功率: {success_rate:.1f}%")
        logger.info(f"  总耗时: {duration:.1f}s")
        logger.info(f"  错误数量: {len(stats['errors'])}")
        logger.info("="*60)


def run_quick_experiment(
    models: Optional[List[str]] = None,
    test_documents: int = 5,
    repetitions: int = 2
) -> Dict[str, Any]:
    """
    运行快速测试实验
    
    Args:
        models: 模型名称列表
        test_documents: 测试文档数量
        repetitions: 重复次数
    
    Returns:
        实验结果
    """
    logger.info("运行快速实验...")
    
    # 转换模型名称为枚举
    model_enums = []
    if models:
        for model_name in models:
            try:
                model_enums.append(LLMModel(model_name))
            except ValueError:
                logger.warning(f"未知模型: {model_name}")
    else:
        model_enums = [LLMModel.DEEPSEEK_R1_DISTILL_QWEN_32B_ALIYUN]
    
    # 创建实验执行器并运行
    runner = ExperimentRunner(models=model_enums)
    return runner.run_experiment(
        repetitions=repetitions,
        test_documents=test_documents,
        save_results=True
    )


def run_full_experiment(
    models: Optional[List[str]] = None,
    repetitions: int = 5
) -> Dict[str, Any]:
    """
    运行完整实验
    
    Args:
        models: 模型名称列表
        repetitions: 重复次数
    
    Returns:
        实验结果
    """
    logger.info("运行完整实验...")
    
    # 转换模型名称为枚举
    model_enums = []
    if models:
        for model_name in models:
            try:
                model_enums.append(LLMModel(model_name))
            except ValueError:
                logger.warning(f"未知模型: {model_name}")
    else:
        # 使用默认模型集合
        model_enums = [
            LLMModel.DEEPSEEK_R1_DISTILL_QWEN_32B_ALIYUN,
            LLMModel.DEEPSEEK_V3_OFFICIAL,
            LLMModel.QWEN3_32B,
            LLMModel.GPT_35_TURBO_OPENROUTER
        ]
    
    # 创建实验执行器并运行
    runner = ExperimentRunner(models=model_enums)
    return runner.run_experiment(
        repetitions=repetitions,
        save_results=True
    )